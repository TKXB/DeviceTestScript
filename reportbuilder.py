# coding:utf-8
# 读取已有的word文档

from docx import Document
import time
import configparser

config_raw = configparser.RawConfigParser()
config_raw.read("./config.cfg")


class GenerateReport(object):
    def __init__(self, device):
        self.deviceinfo = device
        self.PortList = device.PortList
        self.HostName = device.hostname

    def run(self):
        # 加载模板
        document = Document(config_raw.get('DEFAULT', 'WORD_TEMPLATE_PATH'))
        # 读取文档中的所有段落的列表
        tables = document.tables
        # 遍历table，并将所有单元格内容写入文件中
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text == "${date}":
                        cell.text = time.strftime("%Y-%m-%d", time.localtime())
                    if cell.text == "${PortList}":
                        cell.text = "、".join('%s' % port for port in self.PortList)
                    if cell.text == "${HostName}":
                        cell.text = self.HostName
                    if cell.text == "${IsRtspVul}":
                        if self.deviceinfo.isrtspvul:
                            cell.text = "未启用身份认证"
                        else:
                            cell.text = "已启用身份认证"
        if self.HostName != '':
            document.save(self.HostName + ".docx")
        else:
            document.save("device.docx")
        print("报告生成完毕！\n")


# def main():
#
#     # 创建文档对象
#     document = Document('/home/tk/Documents/python.docx')
#
#     # 读取文档中所有的段落列表
#     ps = document.paragraphs
#     # 每个段落有两个属性：style和text
#     ps_detail = [(x.text, x.style.name) for x in ps]
#     with open('out.tmp', 'w+') as fout:
#         fout.write('')
#     # 读取段落并写入一个文件
#     with open('out.tmp', 'a+') as fout:
#         for p in ps_detail:
#             fout.write(p[0] + '\t' + p[1] + '\n\n')
#
#     # 读取文档中的所有段落的列表
#     tables = document.tables
#     # 遍历table，并将所有单元格内容写入文件中
#     with open('out.tmp', 'a+') as fout:
#         for table in tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     if cell.text == "${date}":
#                         cell.text = time.strftime("%Y-%m-%d", time.localtime())
#                     if cell.text == "${PortList}":
#                         global PortList
#                         cell.text = "、".join(PortList)
#
#                     fout.write(cell.text + '\t')
#                 fout.write('\n')
#
#     document.save("/home/tk/Documents/save.docx")
#
#
# if __name__ == '__main__':
#     main()



